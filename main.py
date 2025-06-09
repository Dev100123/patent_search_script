import os
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from collections import Counter
import streamlit as st

# Load .env file
load_dotenv()
SERPAPI_KEY = os.getenv("API")

def fetch_full_patent_text(url, short_title, short_summary):
    try:
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")

        title_tag = soup.find("meta", {"name": "citation_title"})
        abstract_tag = soup.find("meta", {"name": "description"})

        title = title_tag["content"].strip() if title_tag else short_title
        summary = abstract_tag["content"].strip() if abstract_tag else short_summary

        return title, summary
    except Exception:
        return short_title, short_summary

def search_patents(query: str, top_n: int = 10):
    if not SERPAPI_KEY:
        return {"error": "SerpAPI API key not found. Please set it in your .env file."}

    url = "https://serpapi.com/search"
    params = {
        "engine": "google_patents",
        "q": query,
        "api_key": SERPAPI_KEY,
        "num": top_n,
        "output": "json",
        "no_cache": "true",
    }

    response = requests.get(url, params=params)
    if response.status_code != 200:
        return {"error": f"Error from SerpAPI: {response.status_code}"}

    data = response.json()
    results = data.get("organic_results", [])

    patents = []
    assignees = []
    inventors = []

    for result in results[:top_n]:
        link = result.get("patent_link", "")
        short_title = result.get("title", "No title found")
        short_summary = result.get("snippet", "No summary found")

        full_title, full_summary = fetch_full_patent_text(link, short_title, short_summary)

        patent = {
            "title": full_title,
            "summary": full_summary,
            "publication_date": result.get("publication_date", "Unknown"),
            "inventor": result.get("inventor", "Unknown"),
            "assignee": result.get("assignee", "Unknown"),
            "patent_link": link,
            "pdf": result.get("pdf", "")
        }

        if patent["assignee"] != "Unknown":
            assignees.append(patent["assignee"])
        if patent["inventor"] != "Unknown":
            inventors.append(patent["inventor"])

        patents.append(patent)

    summary = {
        "query": query,
        "total": len(patents),
        "top_assignees": Counter(assignees).most_common(3),
        "top_inventors": Counter(inventors).most_common(3),
    }

    return {"summary": summary, "results": patents}

def create_word_doc(data):
    doc = Document()
    doc.add_heading("Patent Search Report", level=1)

    doc.add_paragraph(f"Query: {data['summary']['query']}")
    doc.add_paragraph(f"Total Patents Found: {data['summary']['total']}")

    doc.add_heading("Top Assignees", level=2)
    for name, count in data["summary"]["top_assignees"]:
        doc.add_paragraph(f"{name} ({count})", style="List Bullet")

    doc.add_heading("Top Inventors", level=2)
    for name, count in data["summary"]["top_inventors"]:
        doc.add_paragraph(f"{name} ({count})", style="List Bullet")

    doc.add_heading("Patent Details", level=2)
    for i, p in enumerate(data["results"], 1):
        doc.add_paragraph(f"{i}. {p['title']}", style="List Number")
        doc.add_paragraph(f"Summary: {p['summary']}")
        doc.add_paragraph(f"Publication Date: {p['publication_date']}")
        doc.add_paragraph(f"Inventor: {p['inventor']}")
        doc.add_paragraph(f"Assignee: {p['assignee']}")
        doc.add_paragraph(f"Patent Link: {p['patent_link']}")
        doc.add_paragraph(f"PDF Link: {p['pdf']}")
        doc.add_paragraph("")

    return doc

def main():
    st.set_page_config(
        page_title="Patent Search",
        page_icon="🔎",
        layout="centered"
    )
    
    st.markdown(
        "<h1 style='text-align: center; color: #4B8BBE;'>🔎 Patent Search</h1>",
        unsafe_allow_html=True,
    )
    st.write("---")

    query = st.text_input(
        "🔍 Enter your patent search query:",
        placeholder="E.g. AI for manufacturing in the domain of shoes",
        help="Type your query to search relevant patents."
    )

    top_n = st.number_input(
        "⚙️ Number of patents to fetch (min 11, max 50):",
        min_value=11,
        max_value=50,
        value=11,
        step=1,
        help="Select how many patents to retrieve (must be greater than 10)."
    )

    st.write("")  # Spacer

    col1, col2 = st.columns([1, 3])
    with col1:
        search_btn = st.button("🔎 Search Patents", use_container_width=True)
    with col2:
        st.write("")  # For spacing/alignment

    if search_btn:
        if not query.strip():
            st.error("⚠️ Please enter a query to search.")
            return

        with st.spinner("⏳ Searching patents... Please wait."):
            result = search_patents(query, top_n)

        if "error" in result:
            st.error(f"❌ {result['error']}")
            return

        # Show success message
        st.success(f"✅ Found {result['summary']['total']} patents for your query!")

        # (The rest of your display code goes here, e.g. showing summary, details, download button)


        # Display summary
        summary = result["summary"]
        st.subheader("Summary")
        st.write(f"**Query:** {summary['query']}")
        st.write(f"**Total Patents Found:** {summary['total']}")

        st.write("**Top Assignees:**")
        for name, count in summary["top_assignees"]:
            st.write(f"- {name} ({count})")

        st.write("**Top Inventors:**")
        for name, count in summary["top_inventors"]:
            st.write(f"- {name} ({count})")

        # Display patent details
        st.subheader("Patent Details")
        for i, patent in enumerate(result["results"], 1):
            st.markdown(f"**{i}. {patent['title']}**")
            st.write(f"Summary: {patent['summary']}")
            st.write(f"Publication Date: {patent['publication_date']}")
            st.write(f"Inventor: {patent['inventor']}")
            st.write(f"Assignee: {patent['assignee']}")
            st.write(f"[Patent Link]({patent['patent_link']})")
            if patent['pdf']:
                st.write(f"[PDF Link]({patent['pdf']})")
            st.write("---")

        # Create docx and provide download button
        doc = create_word_doc(result)
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Download Patent Report (Word)",
            data=buffer,
            file_name="patent_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
